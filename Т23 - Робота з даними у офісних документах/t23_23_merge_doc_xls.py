#T23_23
# Злиття даних з файлів MS Word, MS Excel за шаблоном MS Word
# Імена полів, що заповнюються даними, мають бути взяті у фігурні дужки { }
# Файли, з яких треба брати дані, вказують у конфігураційному файлі

from docx import Document
import os

from t23_21_mergesource import *
from T21.t21_21_config_dict import *
                 
class Merger:
    '''Злиття даних у документ MS Word.

       Злиття здійснюється згідно з шаблоном (документом MS Word).
       У шаблоні поля для злиття беруться у фігурні дужки { }.
       Дані для кожного поля отримують з відповідного джерела даних, яким є
       стовпчик таблиці у документі Word або стовпчик аркуша робочої книги Excel.
       self.indoc - об'єкт Document з файлу-шаблону
       self.mergesrc - об'єкт класу MergeSource - містить джерела даних та
                       повертає записи з даними злиття
       self.outfile - ім'я файлу-результату
       self.outdoc - об'єкт Document файлу-результату

    '''
    def __init__(self, template, param_files, leadparam, outfile = None):
        '''Конструктор.

           Здійснює під'єднання до джерел даних.
           param_files - словник, що містить імена параметрів (полів)
           та імена файлів, де розташовано відповідні дані.
           leadparam - провідний параметр.
           Кількість записів даних розраховується за цим параметром.
           template - ім'я файлу-шаблону
           outfile - ім'я файлу-результату. Якщо не вказано, то ім'я будується за template
        '''
        self.indoc = Document(template)
        if not leadparam in param_files: # треба, щоб провідний параметр обов'язково був
            raise KeyError
        # зв'язати з джерелами даних, підготуватися до злиття
        self.mergesrc = MergeSource(param_files, leadparam)
        if not outfile:
            base, ext = os.path.splitext(template) # будуємо ім'я файлу-результату
            outfile = base + '_merged.docx'
        self.outfile = outfile
        self.outdoc = Document()
    
    def merge(self):
        '''Здійснити злиття.'''
        for record in self.mergesrc:
            self._process_template(record)
            self.outdoc.add_page_break() # додати розрив сторінки
        self.outdoc.save(self.outfile)   # зберегти файл результату
    
    def _process_template(self, record):
        '''Обробити файл шаблону для одного запису record.'''
        for paragraph in self.indoc.paragraphs:
            self._process_paragraph(paragraph, record)
            
    def _process_paragraph(self, paragraph, record):
        '''Обробити один параграф paragraph для одного запису record.'''
        # додати порожній параграф до документу-результату
        new_para = self.outdoc.add_paragraph(style=paragraph.style)
        para_format_copy(paragraph, new_para)
        for run in paragraph.runs:
            self._process_run(run, record, new_para)

    def _process_run(self, run, record, new_para):
        '''Обробити один потік тексту run для одного запису record.

           Результат записати у параграф new_para
        '''
        text = run.text
        while True:
            pos, field = self._get_field_pos(text, record) # чи є мітка поля даних
            if pos == -1: break
            if pos > 0: # якщо до мітки є текст, додаємо його до результату
                new_run = new_para.add_run(text[:pos], run.style)
                run_format_copy(run, new_run)
            text = text[pos + len(field) + 2:] # виключаємо доданий текст та мітку поля
            srcitem = record[field] # отримуємо об'єкт з даними SourceItem
            if srcitem.type == 'excel':
                # для Excel усе просто
                new_run = new_para.add_run(str(srcitem.obj.value), run.style)
                run_format_copy(run, new_run)
            else:
                # У клітинці таблиці Word може бути декілька параграфів.
                # Їх усі треба додати. Але додаємо як текстові потоки,
                # а між параграфами вставляємо розрив рядка
                num_para = len(srcitem.obj.paragraphs)
                for i, t_para in enumerate(srcitem.obj.paragraphs):
                    for t_run in t_para.runs:
                        new_run = new_para.add_run(t_run.text, t_run.style)
                        run_format_copy(t_run, new_run)
                    if i < num_para - 1: # якщо параграф не останній, додаємо розрив рядка
                        new_run.add_break()
        # додаємо до результату текст після останньої мітки поля
        # якщо міток немає, то тут додається весь текст
        new_run = new_para.add_run(text, run.style)
        run_format_copy(run, new_run)
        
                        
    def _get_field_pos(self, text, record):
        '''Знайти позицію поля даних з словника record у тексті text.

           Якщо не знайдено, повертає -1.
        '''
        pos = -1
        field = None
        for name in record:
            s = '{' + name + '}' # формуємо рядок мітки
            pos = text.find(s)
            if pos >= 0:
                field = name
                break
        return pos, field
                
def para_format_copy(p1, p2):
    '''Копіює формат параграфа p1 у p2'''
    pf1 = p1.paragraph_format # об'єкт ParagraphFormat
    pf2 = p2.paragraph_format
    pf2.alignment = pf1.alignment
    pf2.line_spacing = pf1.line_spacing
    pf2.first_line_indent = pf1.first_line_indent
    pf2.keep_together = pf1.keep_together
    pf2.keep_with_next = pf1.keep_with_next
    pf2.left_indent = pf1.left_indent
    pf2.line_spacing_rule = pf1.line_spacing_rule
    pf2.page_break_before = pf1.page_break_before
    pf2.right_indent = pf1.right_indent
    pf2.space_after = pf1.space_after
    pf2.space_before = pf1.space_before
    pf2.widow_control = pf1.widow_control
    
def run_format_copy(r1, r2):
    '''Копіює формат потоку символів r1 у r2'''
    r2.bold = r1.bold
    r2.italic = r1.italic
    r2.underline = r1.underline
    rf1 = r1.font   # об'єкт Font
    rf2 = r2.font
    rf2.all_caps = rf1.all_caps
    rf2.color.rgb = rf1.color.rgb # об'єкт ColorFormat
    rf2.name = rf1.name
    rf2.size = rf1.size
    rf2.small_caps = rf1.small_caps
    rf2.strike = rf1.strike
    rf2.subscript = rf1.subscript
    rf2.superscript = rf1.superscript


if __name__ == '__main__':
    import sys
    if len(sys.argv) == 1:              # якщо немає параметрів, ставимо ім'я за угодою
        config = "config.txt"
    else:
        config = sys.argv[1]         # 1 параметр
    conf = ConfigDict(config)
    params = conf.getconfig()
    # окремо виділяємо leadparam, template, outfile
    leadparam = params['LeadParam']
    del params['LeadParam']
    template = params['Template']
    del params['Template']
    outfile = params.get('OutFile')
    if outfile:
        del params['OutFile']

    merger = Merger(template, params, leadparam, outfile)
    merger.merge()
    

    
    


    
